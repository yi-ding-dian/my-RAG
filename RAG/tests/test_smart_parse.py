"""智能解析引导：文档画像分析接口测试

覆盖：
- 权限：普通用户 403；dept_admin 本部门 200 / 跨部门 403；文档不存在 404
- 文本提取：txt/md 直读（utf-8/GBK）、docx 真实提取（python-docx）、
  pdf 复用 _extract_plain（mock）；未解析（uploaded）文档也能分析
- 标题结构：# 标题/包裹式/前导符号式（_iter_headings）+ 编号式补充
  （第X章/第X节/X.X/一、）
- 篇幅：总字符数 + 2 万阈值实时读（改配置即时生效）+ over_threshold
- QA 格式：复用 analyze_qa_format，占比 >=50% 判定
- 指代密集度：规则统计按字数归一化 → low/mid/high
- 推荐矩阵：QA → qa；有标题 → parent_child；无标题 ≤ 阈值 → naive+上下文
  检索推荐；无标题 > 阈值 → naive+不建议上下文检索；agentic 可选
- 容错：源文件缺失 → 部分画像 + warning（接口仍 200）
"""
from __future__ import annotations

import io

import pytest

from conftest import admin_headers_of, create_kb, upload_doc
from backend.services import settings_service as ss

# ==================== 解析器探测 mock（画像接口同款） ====================

_PARSER_PROBE_ALL_OK = {
    "mineru": {"available": True, "reason": ""},
    "deepdoc": {"available": True, "reason": ""},
    "plain": {"available": True, "reason": ""},
}

_PARSER_PROBE_NO_EXTERNAL = {
    "mineru": {"available": False, "reason": "连接失败"},
    "deepdoc": {"available": False, "reason": "连接失败"},
    "plain": {"available": True, "reason": ""},
}

_PARSER_PROBE_DEEPDOC_ONLY = {
    "mineru": {"available": False, "reason": "连接失败"},
    "deepdoc": {"available": True, "reason": ""},
    "plain": {"available": True, "reason": ""},
}


@pytest.fixture(autouse=True)
def _mock_probe(monkeypatch):
    """默认 mock smart_parse 的解析器探测为全部可用（离线测试不连真实服务）"""
    state = {"probe": _PARSER_PROBE_ALL_OK}

    async def _probe(cfg=None, **kw):
        return state["probe"]

    monkeypatch.setattr("backend.routers.smart_parse.probe_parsers", _probe)
    return state


def _set_threshold(value: int):
    """直接改活跃档案阈值并应用全局配置（模拟超管改配置，即时生效）"""
    svc = ss.get_settings_service()
    p = svc.get_active()
    p.setdefault("contextual_retrieval", {})["max_full_doc_chars"] = value
    svc._profiles[p["id"]] = p
    svc._save()
    svc._apply_active()


def _analyze(client, kb_id, doc_id, headers=None):
    """调用画像接口，断言 200 并返回 JSON"""
    hdrs = headers if headers is not None else admin_headers_of(client)
    resp = client.get(f"/api/kbs/{kb_id}/documents/{doc_id}/analyze",
                      headers=hdrs)
    assert resp.status_code == 200, resp.text
    return resp.json()


# ==================== 权限 ====================

class TestPermission:

    def test_user_403(self, client, admin_headers, user_headers):
        """普通用户无 can_manage_kb → 403"""
        kb = create_kb(client)
        doc = upload_doc(client, kb["id"])
        resp = client.get(
            f"/api/kbs/{kb['id']}/documents/{doc['id']}/analyze",
            headers=user_headers)
        assert resp.status_code == 403

    def test_dept_admin_own_kb_ok(self, client, dept_admin_headers):
        """dept_admin 本部门知识库 → 200"""
        kb = create_kb(client, headers=dept_admin_headers)
        doc = upload_doc(client, kb["id"], headers=dept_admin_headers)
        data = _analyze(client, kb["id"], doc["id"],
                        headers=dept_admin_headers)
        assert data["doc_id"] == doc["id"]

    def test_dept_admin_cross_kb_403(self, client, dept_admin_headers,
                                     user_headers):
        """dept_admin 访问其他部门知识库 → 403"""
        kb = create_kb(client)  # admin 建的库（无部门归属）
        doc = upload_doc(client, kb["id"])
        resp = client.get(
            f"/api/kbs/{kb['id']}/documents/{doc['id']}/analyze",
            headers=dept_admin_headers)
        assert resp.status_code == 403

    def test_doc_not_found_404(self, client, admin_headers):
        kb = create_kb(client)
        resp = client.get(
            f"/api/kbs/{kb['id']}/documents/not-exist/analyze",
            headers=admin_headers)
        assert resp.status_code == 404


# ==================== 文本提取 ====================

class TestTextExtraction:

    def test_txt_utf8(self, client, admin_headers):
        """txt 直读（utf-8）"""
        kb = create_kb(client)
        content = "# 标题一\n\n这是正文第一段。\n\n这是正文第二段。"
        doc = upload_doc(client, kb["id"], filename="a.txt",
                         content=content)
        data = _analyze(client, kb["id"], doc["id"])
        assert data["extracted"] is True
        assert data["extract_warning"] is None
        assert data["length"]["doc_chars"] == len(content)
        assert data["file_type"] == "txt"

    def test_txt_gbk(self, client, admin_headers):
        """GBK 编码 txt 回退解码（gbk）"""
        kb = create_kb(client)
        content = "第一章 概述\n\n这是中文正文。".encode("gbk")
        doc = upload_doc(client, kb["id"], filename="b.txt", content=content)
        data = _analyze(client, kb["id"], doc["id"])
        assert data["extracted"] is True
        assert data["length"]["doc_chars"] > 0
        assert data["structure"]["numbered_headings"] >= 1

    def test_md_plain(self, client, admin_headers):
        """md 直读"""
        kb = create_kb(client)
        doc = upload_doc(client, kb["id"], filename="c.md")  # SAMPLE_TEXT
        data = _analyze(client, kb["id"], doc["id"])
        assert data["extracted"] is True
        assert data["file_type"] == "md"

    def test_docx_real_extract(self, client, admin_headers):
        """docx 用 python-docx 真实提取（标题 + 段落 + 表格）"""
        import docx as docx_mod
        buf = io.BytesIO()
        d = docx_mod.Document()
        d.add_heading("报告标题", level=1)
        d.add_paragraph("这是正文第一段。")
        d.add_paragraph("这是正文第二段。")
        table = d.add_table(rows=1, cols=2)
        table.rows[0].cells[0].text = "列A"
        table.rows[0].cells[1].text = "列B"
        d.save(buf)
        kb = create_kb(client)
        doc = upload_doc(client, kb["id"], filename="d.docx",
                         content=buf.getvalue(),
                         mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        data = _analyze(client, kb["id"], doc["id"])
        assert data["extracted"] is True
        # 提取内容含标题段落与表格行
        assert data["length"]["doc_chars"] > 0
        assert data["length"]["paragraphs"] >= 3
        # 标题结构：docx 标题段落非 # 语法，正文段落不含编号 → 无结构（合理）
        assert data["structure"]["heading_count"] == 0

    def test_pdf_via_extract_plain(self, client, admin_headers, monkeypatch):
        """pdf 复用 parser_client._extract_plain（mock 提取，验证接口链路）"""
        import backend.services.parser_client as pc_mod

        def _fake_extract(self, file_path, file_type):
            return "# PDF 标题\n\nPDF 正文内容。\n\n又一段正文。"

        monkeypatch.setattr(pc_mod.ParserClient, "_extract_plain",
                            _fake_extract)
        kb = create_kb(client)
        doc = upload_doc(client, kb["id"], filename="e.pdf",
                         content=b"%PDF-1.4 fake", mime="application/pdf")
        data = _analyze(client, kb["id"], doc["id"])
        assert data["extracted"] is True
        assert data["file_type"] == "pdf"
        assert data["structure"]["has_headings"] is True

    def test_unparsed_doc_analyze(self, client, admin_headers):
        """未解析（uploaded）文档即可画像分析"""
        kb = create_kb(client)
        doc = upload_doc(client, kb["id"])
        assert doc["status"] == "uploaded"
        data = _analyze(client, kb["id"], doc["id"])
        assert data["extracted"] is True


# ==================== 标题结构 ====================

class TestStructure:

    def test_markdown_headings(self, client, admin_headers):
        kb = create_kb(client)
        doc = upload_doc(client, kb["id"])  # SAMPLE_TEXT（# + ## 三级标题）
        data = _analyze(client, kb["id"], doc["id"])
        s = data["structure"]
        assert s["has_headings"] is True
        assert s["heading_count"] == 3
        assert "Python 简介" in s["examples"]

    def test_numbered_headings(self, client, admin_headers):
        content = (
            "第一章 总则\n\n这是正文。\n\n"
            "第一条 为了规范……\n\n"
            "1.1 概述\n\n正文内容。\n\n"
            "一、背景\n\n背景说明。")
        kb = create_kb(client)
        doc = upload_doc(client, kb["id"], filename="编号.txt",
                         content=content)
        data = _analyze(client, kb["id"], doc["id"])
        s = data["structure"]
        assert s["has_headings"] is True
        assert s["numbered_headings"] == 4

    def test_no_headings(self, client, admin_headers):
        content = "这是一段没有标题的正文。\n\n这是另一段正文内容。"
        kb = create_kb(client)
        doc = upload_doc(client, kb["id"], filename="无标题.txt",
                         content=content)
        data = _analyze(client, kb["id"], doc["id"])
        s = data["structure"]
        assert s["has_headings"] is False
        assert s["heading_count"] == 0


# ==================== 篇幅阈值（实时读配置） ====================

class TestLengthThreshold:

    def test_default_threshold_20000(self, client, admin_headers):
        kb = create_kb(client)
        doc = upload_doc(client, kb["id"])
        data = _analyze(client, kb["id"], doc["id"])
        assert data["length"]["threshold_chars"] == 20000
        assert data["length"]["over_threshold"] is False
        assert "字" in data["length"]["threshold_label"]

    def test_threshold_realtime_read(self, client, admin_headers):
        """阈值实时读：改配置（无需重启）→ 接口返回新阈值"""
        kb = create_kb(client)
        doc = upload_doc(client, kb["id"])
        _set_threshold(5000)
        data = _analyze(client, kb["id"], doc["id"])
        assert data["length"]["threshold_chars"] == 5000

    def test_over_threshold_flag(self, client, admin_headers):
        _set_threshold(100)
        kb = create_kb(client)
        doc = upload_doc(client, kb["id"])
        data = _analyze(client, kb["id"], doc["id"])
        assert data["length"]["over_threshold"] is True

    def test_chars_label_format(self, client, admin_headers):
        kb = create_kb(client)
        # 10000 字文本 → "1.0 万字"
        doc = upload_doc(client, kb["id"], filename="长文.txt",
                         content="甲" * 10000)
        data = _analyze(client, kb["id"], doc["id"])
        assert data["length"]["doc_label"] == "1.0 万字"


# ==================== QA 格式 ====================

class TestQa:

    def test_qa_detect(self, client, admin_headers):
        content = (
            "问：什么是 Python？\n\n"
            "答：Python 是一种高级编程语言。\n\n"
            "问：Python 有哪些用途？\n\n"
            "答：Web 开发、数据分析、人工智能等。")
        kb = create_kb(client)
        doc = upload_doc(client, kb["id"], filename="qa.txt", content=content)
        data = _analyze(client, kb["id"], doc["id"])
        q = data["qa"]
        assert q["is_qa"] is True
        assert q["qa_pairs"] == 2
        assert q["ratio"] >= 0.5

    def test_non_qa_detect(self, client, admin_headers):
        kb = create_kb(client)
        doc = upload_doc(client, kb["id"])  # SAMPLE_TEXT 非 QA
        data = _analyze(client, kb["id"], doc["id"])
        assert data["qa"]["is_qa"] is False
        assert data["qa"]["qa_pairs"] == 0


# ==================== 指代密集度 ====================

class TestReferenceDensity:

    def test_dense(self, client, admin_headers):
        para = ("该方案的核心是将其应用于该系统，它显著提升了此流程的效率，"
                "上述做法与其设计初衷一致，前者优于后者，如上所述。")
        content = "\n\n".join([para] * 10)
        kb = create_kb(client)
        doc = upload_doc(client, kb["id"], filename="指代.txt",
                         content=content)
        data = _analyze(client, kb["id"], doc["id"])
        d = data["reference_density"]
        assert d["count"] > 0
        assert d["level"] in ("low", "mid", "high")
        assert d["level_label"]

    def test_sparse(self, client, admin_headers):
        content = "这是第一句。\n\n这是第二句。\n\n这是第三句。"
        kb = create_kb(client)
        doc = upload_doc(client, kb["id"], filename="稀疏.txt",
                         content=content)
        data = _analyze(client, kb["id"], doc["id"])
        d = data["reference_density"]
        assert d["count"] == 0
        assert d["level"] == "low"


# ==================== 推荐矩阵 ====================

class TestRecommendations:

    def test_has_headings_parent_child(self, client, admin_headers):
        kb = create_kb(client)
        doc = upload_doc(client, kb["id"])  # SAMPLE_TEXT 有标题
        data = _analyze(client, kb["id"], doc["id"])
        rec = data["recommendations"]
        assert rec["chunk_method"]["method"] == "parent_child"
        assert rec["chunk_method"]["recommended"] is True
        assert "标题" in rec["chunk_method"]["reason"]
        # 备选：title + agentic 可选
        methods = [a["method"] for a in rec["alternatives"]]
        assert "title" in methods and "agentic" in methods
        # 有结构推荐包含父标题
        assert rec["enable_heading_in_content"] is True
        # 有结构不主推上下文检索
        assert rec["contextual_retrieval"]["recommended"] is False

    def test_no_headings_under_threshold(self, client, admin_headers):
        content = "无标题短文第一段。\n\n无标题短文第二段。"
        kb = create_kb(client)
        doc = upload_doc(client, kb["id"], filename="无标题.txt",
                         content=content)
        data = _analyze(client, kb["id"], doc["id"])
        rec = data["recommendations"]
        assert rec["chunk_method"]["method"] == "naive"
        assert rec["contextual_retrieval"]["recommended"] is True
        assert "上下文检索" in rec["contextual_retrieval"]["reason"]

    def test_no_headings_over_threshold(self, client, admin_headers):
        _set_threshold(200)
        kb = create_kb(client)
        doc = upload_doc(client, kb["id"], filename="长文.txt",
                         content="无标题长文内容。" * 100)
        data = _analyze(client, kb["id"], doc["id"])
        rec = data["recommendations"]
        assert rec["chunk_method"]["method"] == "naive"
        assert rec["contextual_retrieval"]["recommended"] is False
        assert "不建议" in rec["contextual_retrieval"]["reason"]

    def test_qa_document(self, client, admin_headers):
        content = (
            "问：问题一？\n\n答：答案一。\n\n问：问题二？\n\n答：答案二。")
        kb = create_kb(client)
        doc = upload_doc(client, kb["id"], filename="qa.txt", content=content)
        data = _analyze(client, kb["id"], doc["id"])
        rec = data["recommendations"]
        assert rec["chunk_method"]["method"] == "qa"
        assert rec["chunk_method"]["recommended"] is True
        assert "问答对" in rec["chunk_method"]["reason"]


# ==================== 引擎建议 ====================

class TestEngineSuggestion:

    def test_txt_plain(self, client, admin_headers):
        kb = create_kb(client)
        doc = upload_doc(client, kb["id"], filename="a.txt",
                         content="纯文本")
        data = _analyze(client, kb["id"], doc["id"])
        assert data["engine_suggestion"]["suggested"] == "plain"

    def test_pdf_mineru_ok(self, client, admin_headers, _mock_probe):
        kb = create_kb(client)
        doc = upload_doc(client, kb["id"], filename="e.pdf",
                         content=b"%PDF-1.4 fake", mime="application/pdf")
        data = _analyze(client, kb["id"], doc["id"])
        assert data["engine_suggestion"]["suggested"] == "mineru"

    def test_pdf_mineru_down_deepdoc_up(self, client, admin_headers,
                                        _mock_probe):
        _mock_probe["probe"] = _PARSER_PROBE_DEEPDOC_ONLY
        kb = create_kb(client)
        doc = upload_doc(client, kb["id"], filename="e.pdf",
                         content=b"%PDF-1.4 fake", mime="application/pdf")
        data = _analyze(client, kb["id"], doc["id"])
        assert data["engine_suggestion"]["suggested"] == "deepdoc"

    def test_pdf_all_down_plain(self, client, admin_headers, _mock_probe):
        _mock_probe["probe"] = _PARSER_PROBE_NO_EXTERNAL
        kb = create_kb(client)
        doc = upload_doc(client, kb["id"], filename="e.pdf",
                         content=b"%PDF-1.4 fake", mime="application/pdf")
        data = _analyze(client, kb["id"], doc["id"])
        assert data["engine_suggestion"]["suggested"] == "plain"


# ==================== 提取失败容错 ====================

class TestFailureTolerance:

    def test_source_file_missing(self, client, admin_headers):
        """源文件缺失 → 部分画像 + warning，接口仍 200"""
        from backend.config import UPLOAD_DIR
        kb = create_kb(client)
        doc = upload_doc(client, kb["id"], filename="a.txt",
                         content="正文内容")
        (UPLOAD_DIR / doc["name"]).unlink()
        data = _analyze(client, kb["id"], doc["id"])
        assert data["extracted"] is False
        assert data["extract_warning"] is not None
        assert data["warnings"], "应有 warning"
        # 画像字段仍存在（部分画像）
        assert data["length"]["doc_chars"] == 0
        assert data["structure"]["has_headings"] is False
        assert data["qa"]["is_qa"] is False
        assert data["recommendations"]["chunk_method"]["method"] == "naive"

    def test_unknown_file_type(self, client, admin_headers):
        """不支持的类型（如 url 导入文档）→ 提取失败但整体 200"""
        from backend.services.document_service import get_document_service
        kb = create_kb(client)
        doc = upload_doc(client, kb["id"], filename="x.txt",
                         content="正文")
        item = get_document_service().get(doc["id"])
        item.file_type = "url"  # 模拟 URL 导入文档
        data = _analyze(client, kb["id"], doc["id"])
        assert data["extracted"] is False
        assert "暂不支持" in data["extract_warning"]
        assert data["warnings"]
