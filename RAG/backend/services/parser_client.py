"""文档解析客户端（解析引擎可选：MinerU 高精度 / 纯文本提取 / 自动）

- parse(file_path, file_type, engine="auto", **parse_opts):
  - engine="auto"（默认）: pdf/docx 先健康探测 MinerU（GET /health），可用则
    POST /file_parse（multipart）取 markdown，不可用或失败降级纯文本提取
  - engine="mineru": 强制走 MinerU，不可用或失败抛异常（上层标记 failed，
    错误提示指向 MinerU 服务地址，可改用自动/纯文本模式）
  - engine="deepdoc": 强制走 DeepDoc（RAGFlow API，见 deepdoc_client.py），
    表格输出为 HTML <table> 可检索；仅 PDF 支持（docx 抛异常）；不可用或
    失败抛异常（上层标记 failed，错误提示指向 DeepDoc 服务地址）
  - engine="plain": 跳过 MinerU，直接 pypdf/python-docx 纯文本提取
  - parse_opts: 解析参数透传（table_enable/formula_enable/return_images/
    lang_list/pages/backend），以 multipart form 字段随 /file_parse 发送；
    mineru-api 未声明的 form 字段 FastAPI 自动忽略（不报错）
- txt/md: 直接读取（UTF-8 优先，失败回退 GBK），不受引擎选择影响
- 返回 (text, images, parse_method)，parse_method ∈ {mineru, plain}
- images 防御式归一化为 [{name, data: bytes|None}]，容忍四种形态：
  ① images:{文件名: "data:image/xxx;base64,..."}（mineru-api v3 dict 形态，
     解码失败跳过该图不阻塞）
  ② images:[{name, data}]（data 可为 base64 字符串或 bytes）
  ③ images:[base64字符串]
  ④ images:[文件名]
  无 images 字段返回 []
- 扫描版 PDF 提取为空 -> 上层标记 failed 并提示"请启动 MinerU"
- 运行时读取 get_active_config()（配置档案即时生效）
"""
from __future__ import annotations

import base64
import logging
import re
from pathlib import Path
from typing import List, Tuple

import httpx

from backend.config import get_active_config
from backend.services.probes import probe_mineru

logger = logging.getLogger(__name__)

# base64 特征：仅含字母数字 +/+/
_B64_RE = re.compile(r"^[A-Za-z0-9+/=]+$")

# ---- MinerU /file_parse 请求 form 字段（parse_opts 透传）----
# OpenAPI 实测（MINERU_API_URL/openapi.json，mineru-api 官方服务）：
# - files（必填，multipart 数组）；lang_list 为**数组**（默认 ['ch']，重复
#   同名 form 字段传多值）；formula_enable/table_enable/return_images 为
#   boolean（form 传 "true"/"false"）；页码参数为 start_page_id/end_page_id
#   （整数，**从 0 开始**，默认 0/99999）；return_images 默认 False（需显式 true）
# - 未声明的 form 字段 FastAPI 自动忽略（不报错），未知参数安全忽略
# 布尔开关字段名（multipart 传 "true"/"false"）
_MINERU_BOOL_FIELDS = ("table_enable", "formula_enable", "return_images",
                       "return_md", "return_middle_json")
# 数组字段：lang_list（ch/en，multipart 重复同名 key 传多值）
_MINERU_LIST_FIELDS = ("lang_list",)
# 字符串透传字段（backend/parse_method 等，mineru-api 顶层 form 契约）
_MINERU_STR_FIELDS = ("backend", "parse_method")


def _build_mineru_form_data(parse_opts: dict) -> dict:
    """parse_opts → /file_parse 的 multipart form 字段

    全部参数为顶层 form 字段（mineru-api v3 OpenAPI 契约，实测确认）：
    - 布尔开关同名透传（table_enable/formula_enable/return_images/
      return_md/return_middle_json，multipart 传 "true"/"false"）
    - lang_list 传数组（OpenAPI 契约 List[str]，httpx 对 list 值自动
      展开为重复同名 form 字段）
    - 字符串字段同名透传（backend/parse_method，调用方传入才出现）
    - pages（[[from,to],...]，1 基）展开为 start_page_id/end_page_id（0 基：
      from-1/to-1）：单组取 [from,to]，多组仅第一组生效（简化：后续组预留）；
      OpenAPI 实测参数名为 start_page_id/end_page_id，非 from_page/to_page
    """
    data: dict = {}
    for key in _MINERU_BOOL_FIELDS:
        if parse_opts.get(key) is not None:
            data[key] = "true" if parse_opts[key] else "false"
    for key in _MINERU_LIST_FIELDS:
        if parse_opts.get(key):
            value = parse_opts[key]
            data[key] = ([str(v) for v in value]
                         if isinstance(value, (list, tuple))
                         else [str(value)])
    for key in _MINERU_STR_FIELDS:
        if parse_opts.get(key):
            data[key] = str(parse_opts[key])
    pages = parse_opts.get("pages")
    if isinstance(pages, list) and pages:
        group = pages[0]  # 多组页码：仅第一组生效（后续组预留）
        if isinstance(group, (list, tuple)) and len(group) == 2:
            data["start_page_id"] = str(max(int(group[0]) - 1, 0))
            data["end_page_id"] = str(max(int(group[1]) - 1, 0))
    return data


def _decode_image_data(data) -> bytes | None:
    """图片 data 字段解码：bytes 直接返回；base64 字符串解码；其余 None"""
    if data is None:
        return None
    if isinstance(data, bytes):
        return data
    if isinstance(data, str):
        s = data.strip()
        if not s:
            return None
        # 去掉 data:image/png;base64, 前缀（MinerU 变体）
        if s.startswith("data:"):
            s = s.split(",", 1)[-1] if "," in s else s
        try:
            return base64.b64decode(s)
        except Exception:
            return None
    return None


def _normalize_images(raw) -> List[dict]:
    """防御式归一化 → [{name, data: bytes|None}]（无 images 字段/异常形态 → []）

    支持 dict 形态（mineru-api v3 实测：{文件名: "data:image/xxx;base64,..."}，
    文件名与 md 里 ![](images/xxx.jpg) 引用的文件名一致）与 list 形态
    （[{name,data}] / [base64] / [文件名]，历史兼容）。dict 分支中
    解码失败的条目跳过（不阻塞整体，与形态③解码失败视作文件名的行为区分）。
    """
    if raw is None:
        return []
    if isinstance(raw, dict):
        result: List[dict] = []
        for name, val in raw.items():
            decoded = _decode_image_data(val)
            if not decoded:  # None 或空字节（base64 仅非法字符时不抛异常返回 b''）
                logger.warning("解析图片解码失败，跳过: %s", name)
                continue
            result.append({"name": str(name), "data": decoded})
        return result
    if not isinstance(raw, list):
        return []
    result: List[dict] = []
    for i, item in enumerate(raw):
        if isinstance(item, dict):
            name = item.get("name") or item.get("filename") or f"image_{i}"
            data = item.get("data") or item.get("content") or item.get("base64")
            result.append({"name": str(name), "data": _decode_image_data(data)})
        elif isinstance(item, str):
            stripped = item.strip()
            if not stripped:
                continue
            # 形态② base64 字符串（可能有 data:image/xxx;base64, 前缀）：
            # 有长度 + 字符特征才尝试解码，解码失败视为形态③文件名
            payload = stripped
            if payload.startswith("data:"):
                payload = payload.split(",", 1)[-1] if "," in payload else payload
            if (len(payload) > 8 and _B64_RE.match(payload)
                    and len(payload) % 4 == 0):
                try:
                    decoded = base64.b64decode(payload)
                    if decoded:
                        result.append({"name": f"image_{i}", "data": decoded})
                        continue
                except Exception:
                    pass
            result.append({"name": stripped, "data": None})
        else:
            continue
    return result


class ParserClient:

    # ---------- 健康探测 ----------

    async def _mineru_available(self, api_url: str) -> bool:
        """MinerU 健康探测（复用 probes 统一探测：<500 可用，3s 超时）"""
        r = await probe_mineru({"url": api_url}, timeout=3.0)
        return bool(r["ok"])

    # ---------- MinerU 解析 ----------

    async def _parse_via_mineru(self, api_url: str, file_path: Path,
                                timeout: float,
                                **parse_opts) -> Tuple[str, List[dict]]:
        """POST /file_parse（multipart），返回 (markdown 文本, images 列表)

        mineru-api v3 契约（MINERU_API_URL 配置，服务名 mineru-api）：
        - multipart 字段名为 files（数组），非 file；
        - 额外 form 字段透传解析配置（_build_mineru_form_data，OpenAPI 实测）：
          table_enable/formula_enable/return_images/return_md/
          return_middle_json/lang_list（数组）/start_page_id/end_page_id
          （页码，0 基）/backend/parse_method；
        - 接口同步等待任务完成（实测 demo.pdf 约 2 分钟），返回
          {status, results: {<文件名去扩展>: {md_content, ...}}}，
          return_images=True 时 results.<名>.images 为 dict：
          {文件名: "data:image/xxx;base64,..."}（文件名与 md 引用一致）；
        同时兼容旧版 {markdown/text/result} 直返结构与 code/data 嵌套
        """
        form_data = _build_mineru_form_data(parse_opts)
        async with httpx.AsyncClient(timeout=timeout) as client:
            with open(file_path, "rb") as f:
                resp = await client.post(
                    f"{api_url.rstrip('/')}/file_parse",
                    files={"files": (file_path.name, f)},
                    data=form_data or None,
                )
            resp.raise_for_status()
            data = resp.json()
        # 兼容不同返回结构: results.<名>.md_content / markdown / text / result / code/data 嵌套
        text = ""
        images_raw = None
        if isinstance(data, dict):
            # mineru-api v3：任务同步完成，结果按原始文件名（去扩展名）列在 results 下
            results = data.get("results")
            if isinstance(results, dict):
                for item in results.values():
                    if not isinstance(item, dict):
                        continue
                    md = (item.get("md_content")
                          or item.get("markdown") or item.get("text"))
                    if isinstance(md, str) and md.strip():
                        text = md
                    if "images" in item:
                        images_raw = item["images"]
                    break  # 单文件解析，取第一个结果
            if not text:
                for key in ("markdown", "text", "result"):
                    if isinstance(data.get(key), str) and data[key].strip():
                        text = data[key]
                        break
            if images_raw is None and "images" in data:
                images_raw = data["images"]
            if isinstance(data.get("code"), (int, float)) and "data" in data:
                inner = data["data"]
                if isinstance(inner, dict):
                    if not text:
                        for key in ("markdown", "text"):
                            if isinstance(inner.get(key), str) and inner[key].strip():
                                text = inner[key]
                                break
                    if images_raw is None and "images" in inner:
                        images_raw = inner["images"]
        images = _normalize_images(images_raw)
        if images:
            logger.info("MinerU 返回 %d 张图片", len(images))
        return text, images

    # ---------- 降级提取 ----------

    def _extract_plain(self, file_path: Path, file_type: str) -> str:
        """pypdf / python-docx 纯文本提取"""
        try:
            if file_type == "pdf":
                from pypdf import PdfReader
                reader = PdfReader(str(file_path))
                pages = []
                for page in reader.pages:
                    pages.append(page.extract_text() or "")
                return "\n\n".join(pages).strip()
            if file_type == "docx":
                import docx  # python-docx
                d = docx.Document(str(file_path))
                paras = [p.text for p in d.paragraphs if p.text.strip()]
                for table in d.tables:
                    for row in table.rows:
                        cells = [c.text.strip() for c in row.cells]
                        if any(cells):
                            paras.append(" | ".join(cells))
                return "\n\n".join(paras).strip()
        except Exception as e:
            logger.warning("降级提取失败 %s: %s", file_path.name, e)
            return ""
        return ""

    # ---------- 强制 DeepDoc（engine=deepdoc） ----------

    async def _parse_via_deepdoc_forced(self, file_path: Path,
                                        file_type: str) -> Tuple[str, List[dict], str]:
        """强制 DeepDoc：通过 ragflow-server（RAGFlow API）解析 PDF

        DeepDoc 核心价值：表格输出为 HTML <table> 文本可检索（vs MinerU
        表格为图片不可检索）；图片不进 content（无图片链路，images 恒 []）。
        仅支持 PDF（docx 不支持 → 抛异常提示改用其他引擎，上层标记 failed）；
        不可用/失败抛 RuntimeError（错误提示含服务地址，上层标记 failed）
        """
        if file_type != "pdf":
            raise RuntimeError(
                "DeepDoc 仅支持 PDF 文档（docx 请改用 MinerU/自动/纯文本模式）")
        from backend.services.deepdoc_client import parse_via_deepdoc
        cfg = get_active_config().deepdoc
        try:
            text, images = await parse_via_deepdoc(file_path, cfg)
        except Exception as e:
            raise RuntimeError(
                f"DeepDoc 解析失败: {e}（{cfg.base_url}），"
                f"请检查 DeepDoc 服务（系统设置）或改用其他引擎") from e
        if not text.strip():
            raise RuntimeError(
                "DeepDoc 解析结果为空（扫描版 PDF 或无文本内容），"
                "请改用自动/纯文本模式")
        return text, images, "deepdoc"

    # ---------- 强制 MinerU（engine=mineru） ----------

    async def _parse_via_mineru_forced(self, cfg, file_path: Path,
                                       file_type: str,
                                       **parse_opts) -> Tuple[str, List[dict], str]:
        """强制 MinerU：不可用/失败抛 RuntimeError（错误提示含服务地址，上层标记 failed）"""
        if not await self._mineru_available(cfg.api_url):
            raise RuntimeError(
                f"MinerU 解析不可用（{cfg.api_url} 无响应），"
                f"请检查 MinerU 服务或改用自动/纯文本模式")
        try:
            text, images = await self._parse_via_mineru(
                cfg.api_url, file_path, cfg.timeout, **parse_opts)
        except Exception as e:
            raise RuntimeError(
                f"MinerU 解析失败: {e}（{cfg.api_url}），"
                f"请检查 MinerU 服务或改用自动/纯文本模式") from e
        if not text.strip():
            raise RuntimeError(
                "MinerU 解析结果为空（扫描版 PDF 或无文本内容），请改用自动/纯文本模式")
        return text, images, "mineru"

    # ---------- 主入口 ----------

    async def parse(self, file_path: Path, file_type: str,
                    engine: str = "auto",
                    **parse_opts) -> Tuple[str, List[dict], str]:
        """解析文件，返回 (text, images, parse_method)

        engine ∈ {auto, mineru, deepdoc, plain}（非法值按 auto 处理，防御式）：
        - auto（默认）: MinerU 可用则用，不可用/失败降级 plain（原有行为）
        - mineru: 强制走 MinerU，不可用/失败抛异常（上层标记 failed）
        - deepdoc: 强制走 DeepDoc（RAGFlow API，表格输出为可检索 HTML），
          仅 PDF 支持，不可用/失败抛异常（上层标记 failed）
        - plain: 跳过 MinerU，直接 pypdf/python-docx 纯文本提取
        - parse_opts: 解析参数透传（table_enable/formula_enable/return_images/
          lang_list/pages，见 _build_mineru_form_data），txt/md 直读不受影响
        txt/md 直读不受引擎影响
        """
        file_type = (file_type or "").lower().lstrip(".")
        # url = URL 网页导入文档（内容即提取出的纯文本 .md，与 txt/md 同路径直读）
        if file_type in ("txt", "md", "markdown", "url"):
            for enc in ("utf-8", "gbk"):
                try:
                    return file_path.read_text(encoding=enc), [], "plain"
                except UnicodeDecodeError:
                    continue
            return file_path.read_text(encoding="utf-8", errors="replace"), [], "plain"

        if file_type in ("pdf", "docx"):
            cfg = get_active_config().mineru
            if engine == "deepdoc":
                # DeepDoc：RAGFlow 解析（表格输出为可检索 HTML）；仅 PDF，
                # docx 不支持 → 抛异常；MinerU 解析参数对其不适用
                return await self._parse_via_deepdoc_forced(
                    file_path, file_type)
            if engine == "mineru":
                return await self._parse_via_mineru_forced(
                    cfg, file_path, file_type, **parse_opts)
            if engine == "plain":
                text = self._extract_plain(file_path, file_type)
                return text, [], "plain"
            if engine != "auto":
                logger.warning("未知解析引擎 %r，按 auto 处理", engine)
            # auto：原有探测 + 降级逻辑
            if await self._mineru_available(cfg.api_url):
                try:
                    text, images = await self._parse_via_mineru(
                        cfg.api_url, file_path, cfg.timeout, **parse_opts)
                    if text.strip():
                        return text, images, "mineru"
                    logger.warning("MinerU 解析结果为空，降级 plain: %s", file_path.name)
                except Exception as e:
                    logger.warning("MinerU 解析失败，降级 plain: %s", e)
            else:
                logger.info("MinerU 不可用(%s)，降级 plain", cfg.api_url)
            text = self._extract_plain(file_path, file_type)
            return text, [], "plain"

        raise ValueError(f"不支持的文件类型: {file_type}")


_parser_client: ParserClient | None = None


def get_parser_client() -> ParserClient:
    global _parser_client
    if _parser_client is None:
        _parser_client = ParserClient()
    return _parser_client
